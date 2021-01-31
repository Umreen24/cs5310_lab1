# importing packages for lab
import os
import numpy as np
import pandas as pd
import math
from pyeeg import bin_power
import matplotlib.pyplot as plt
from docx import Document
from scipy.io import loadmat

# path to acquisition file
os.chdir('/Users/umreenimam/documents/masters/masters_classes/cs_5310/week_2/lab_1')
# setting filename variable
filename = 'Acquisition-15-data.mat'
# loading data into environment via loadmat func from scipy.io
data = loadmat(filename)
# setting eeg variable to the contents of 'data' variable
eeg = data['data']
# transpose array and put into data frame
eeg_transposed = np.transpose(eeg)

##############
# QUESTIONS 1-4
##############
# EEG values mean
# using axis=1 to run function across 34 channels
data_means = np.mean(eeg_transposed, axis=1)

# EEG values standard deviation
# using axis=1 to run function across 34 channels
data_sd = np.std(eeg_transposed, axis=1)

# calculating plot info
means = data_means
mean_diff = means - data_sd
mean_sum = means + data_sd

# creating scatterplot of means
# getting time values for scatterplot x-axis
x_vals = np.arange(len(data_means))/1024

plt.scatter(x_vals, mean_diff, color='salmon', edgecolors='salmon', label='Mean - Stddev')
plt.scatter(x_vals, means, color='midnightblue', edgecolors='midnightblue', label='Means')
plt.scatter(x_vals, mean_sum, color='seagreen', edgecolors='seagreen', label='Mean + Stddev')
plt.title('Scatterplot of EEG Means Over Time')
plt.xlabel('Time in Seconds')
plt.ylabel('EEG Value')
plt.legend()
plt.savefig('figure_one.png')
plt.show()

###############
# QUESTIONS 5-8
###############
# loop variables
alpha_psi = pd.DataFrame([])
band = [0.5, 4, 7, 12, 30]
fs = 1024
sample = 1024

# getting number of columns and rows
# columns
columns = eeg_transposed.shape[1]
# rows
rows = math.floor(eeg_transposed.shape[0] / sample)
# x values for scatterplot of alpha psis
alpha_x_vals = np.arange(rows)

# creating function to run for loop
def get_alphas():
    for x in range(columns):
        bin_powers = []
        for y in range(rows):
            psis = bin_power(eeg_transposed[(y * sample):((y + 1) * sample), x], band, fs)[0][2]
            bin_powers.append(psis)
        alpha_psi[x] = bin_powers
    return alpha_psi

# calling function
get_alphas()

# calculations for means & stddev
alpha_means = np.mean(alpha_psi, axis=1)
alpha_stddev = np.std(alpha_psi, axis=1)

# cals for sums & diffs
alpha_diff = alpha_means - alpha_stddev
alpha_sum = alpha_means + alpha_stddev

plt.scatter(alpha_x_vals, alpha_diff, color='salmon', edgecolors='salmon', label='Mean - Stddev')
plt.scatter(alpha_x_vals, alpha_means, color='midnightblue', edgecolors='midnightblue', label='Means')
plt.scatter(alpha_x_vals, alpha_sum, color='seagreen', edgecolors='seagreen', label='Mean + Stddev')
plt.title('Scatterplot of Alpha PSI Mean Values Over Time')
plt.xlabel('Time in Seconds')
plt.ylabel('Alpha Value')
plt.legend()
plt.savefig('figure_two.png')
plt.show()

##################
# QUESTIONS 9 & 10
##################
# adding figures to Word Document
new_doc = Document()
new_doc.add_picture('figure_one.png')
new_doc.add_picture('figure_two.png')
new_doc.save('Lab1_Figures.docx')

# Exporting dataframe to .csv file
alpha_psi.to_csv('Acquisition-15-alphaPSI.csv')
